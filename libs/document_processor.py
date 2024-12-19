from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement, parse_xml
from docx.text.run import Run
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
import pythoncom
from PIL import Image
import os
import json
import re
import uuid
import asyncio
from openai import AsyncOpenAI

class DocumentProcessor:
    def __init__(self, TEMPLATE_FOLDER, UPLOAD_FOLDER, STATIC_DIR, api_key):
        self.UPLOAD_FOLDER = UPLOAD_FOLDER
        self.STATIC_DIR = STATIC_DIR
        self.TEMPLATE_FOLDER = TEMPLATE_FOLDER
        self.client = AsyncOpenAI(api_key=api_key)

    async def generate_docx(self, generated_sections, image_list=None, is_temp=False, index = None):
        if is_temp:
            output_filename = f"temp_business_plan_{uuid.uuid4()}.docx"
        else:
            output_filename = f"business_plan_{uuid.uuid4()}.docx"
        output_filename = f"test1_business_plan_{uuid.uuid4()}.docx"
        output_path = os.path.join(self.UPLOAD_FOLDER, output_filename)
        
        if index == 1:
            template_path = os.path.join(self.TEMPLATE_FOLDER, 'documents_tag_meating_resection.docx')
        elif index == 2:
            template_path = os.path.join(self.TEMPLATE_FOLDER, 'documents_tag_meating_resection.docx')
        elif index == 3:
            template_path = os.path.join(self.TEMPLATE_FOLDER, 'documents_tag_meating_resection.docx')
        else:
            template_path = os.path.join(self.TEMPLATE_FOLDER, 'documents_tag_meating_resection.docx')

        # path 값만 추출하여 리스트로 만듦
        user_key_image = [item['path'] for item in image_list] if image_list else []

        style_mapping = {
            'Heading 1': '개요 1',
            'Heading 2': '개요 2',
            'Heading 3': '개요 3',
        }

        user_keys = ['SECTION1-A2', 'SECTION1-A3', 'SECTION1-A6', 'SECTION1-A7', 
                    'SECTION2-A2', 'SECTION2-A3', 'SECTION2-A6', 'SECTION2-A7', 
                    'SECTION3-A2', 'SECTION3-A3', 'SECTION3-A6', 'SECTION3-A7',
                    'SECTION3-A10', 'SECTION3-A11','SECTION3-A13', 'SECTION3-A14',
                    'SECTION5-A2', 'SECTION5-A3']

        sub_tags = self.generate_dynamic_tags(user_keys)

        final_texts = {}
        
        try:
            document = Document(template_path)
            # 태그 패턴 수정: 정확한 태그 매칭을 위해
            # pattern = re.compile(r'\b([A-Z][a-z]*)?\d+\b')
            pattern = re.compile(r'\b(?:SECTION\d+-[A-Z]+\d*|T\d+)\b')

            # 문단과 표의 모든 문단을 순회하기 위한 함수
            def iter_paragraphs(parent):
                for paragraph in parent.paragraphs:
                    yield paragraph
                for table in parent.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in iter_paragraphs(cell):
                                yield paragraph

            # 문서의 모든 문단을 순회
            for paragraph in iter_paragraphs(document):
                full_text = ''.join(run.text for run in paragraph.runs).strip()
                match = pattern.fullmatch(full_text)
                if match:
                    tag = match.group()  # 태그 그대로 사용
                    if tag in generated_sections:
                        content = generated_sections[tag]
                        # 기존 런에서 태그 제거
                        for run in paragraph.runs:
                            if tag in run.text:
                                run.text = run.text.replace(tag, "")
                                
                        new_text = full_text.replace(tag, "")
                        for run in paragraph.runs:
                            run.text = ""  # 기존 텍스트 제거

                        if tag.startswith('T'):  # 'T'로 시작하는 태그 처리
                            if content.get('text'):
                                text_to_insert = content['text']
                                run = paragraph.add_run(text_to_insert)  # 텍스트 삽입
                                # 스타일 설정
                                font = run.font
                                font.size = Pt(10)  # 글씨 크기
                                font.name = '맑은 고딕'  # 글씨체 설정
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 가운데 정렬
                                
                                if text_to_insert:
                                    final_texts[tag] = text_to_insert.strip()

                        elif any(url.split('/')[-2] == tag for url in user_key_image):
                            # 이 부분을 따로 처리
                            # user_key_image에서 tag와 일치하는 URL 찾기
                            matching_url = next(
                                (url for url in user_key_image if url.split('/')[-2] == tag),
                                None
                            )
                            if matching_url:
                                # 절대 경로 생성
                                relative_path = matching_url.replace('.jpg', '.docx')

                                if relative_path.startswith('/static'):
                                    relative_path = relative_path[len('/static'):]

                                image_docx_path = os.path.abspath(os.path.join(self.STATIC_DIR, relative_path.strip('/')))

                                # 중복된 'static/static/' 제거
                                if 'static/static/' in image_docx_path:
                                    image_docx_path = image_docx_path.replace('static/static/', 'static/')

                                if os.path.exists(image_docx_path):
                                    image_docx = Document(image_docx_path)
                                    # 템플릿 내용의 XML 가져오기
                                    template_element = image_docx.element.body
                                    copied_element = parse_xml(template_element.xml)
                                    # 현재 문단의 부모 요소와 위치 찾기
                                    parent_element = paragraph._element.getparent()
                                    index = parent_element.index(paragraph._element) + 1
                                    # 태그가 있는 문단 제거
                                    parent_element.remove(paragraph._element)
                                    # 복사된 템플릿 내용을 삽입
                                    for child in reversed(copied_element):
                                        parent_element.insert(index, child)

                        elif tag in sub_tags.keys():
                            # 이미지 삽입
                            if content.get('image'):
                                image_path = content['image']
                                if os.path.exists(image_path):
                                    with Image.open(image_path) as img:
                                        desired_width = Inches(2.8)  # 필요에 따라 조정
                                        aspect_ratio = img.height / img.width
                                        calculated_height = desired_width * aspect_ratio
                                    paragraph.add_run().add_picture(image_path, width=desired_width, height=calculated_height)
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            if content.get('image'):
                                image_path = content.get('image')
                                if os.path.exists(image_path):
                                    with Image.open(image_path) as img:
                                        aspect_ratio = img.height / img.width
                                        calculated_height = desired_width * aspect_ratio
                                    paragraph.add_run().add_picture(image_path, width=desired_width, height=calculated_height)

                            # 텍스트 삽입
                            if content.get('text'):
                                await self.parse_markdown_and_insert(content['text'], paragraph, style_mapping)

                                lines = content['text'].split('\n')
                                final_text = ''
                                for line in lines:
                                    final_text += line + '\n'
                                if final_text:
                                    final_texts[tag] = final_text.strip()
                else:
                    # 'A 2-1' 또는 'A 3-1'을 찾는 로직
                    if full_text in sub_tags.values():
                        main_tag = [k for k, v in sub_tags.items() if v == full_text][0]
                        content = generated_sections.get(main_tag)
                        if content:
                                    
                            # paragraph의 모든 runs 제거
                            for run in paragraph.runs[:]:  # [:] 로 복사본을 만들어 순회
                                paragraph._element.remove(run._element)
                                
                            # 텍스트 삽입
                            if content.get('text'):
                                text_to_insert = content['text']
                                run = paragraph.add_run(f"<{text_to_insert}>")  # <> 사이에 텍스트 삽입
                                # 스타일 설정
                                font = run.font
                                font.size = Pt(10)  # 글씨 크기
                                font.name = '맑은 고딕'  # 글씨체 설정
                                font.color.rgb = RGBColor(255, 255, 255)  # 흰색 텍스트
                                run.bold = True
                            else:
                                # 텍스트가 없으면 이미지 파일명 삽입
                                image_path = content.get('image')
                                if image_path:
                                    image_filename = os.path.basename(image_path)
                                    run = paragraph.add_run(f"<{image_filename}>")
                                    # 이미지 파일명에도 동일한 스타일 적용
                                else:
                                    # 나중에 파일명 말고 해당 로직을 사용하고 싶으시면 if 문 없애고 해당 코드를 상단으로 올리시면 됩니다.
                                    tag_value = sub_tags[main_tag]
                                    text_mapping = {
                                        '<SECTION1-A 2-1>': '<제품/서비스 세부내용>',
                                        '<SECTION1-A 3-1>': '<핵심 기능 설명>',
                                        '<SECTION1-A 6-1>': '<개발 목표>',
                                        '<SECTION1-A 7-1>': '<구현 방안>',
                                        '<SECTION2-A 2-1>': '<시장 분석>',
                                        '<SECTION2-A 3-1>': '<시장 특성>',
                                        '<SECTION2-A 6-1>': '<수요 예측>',
                                        '<SECTION2-A 7-1>': '<공급 현황>',
                                        '<SECTION3-A 2-1>': '<구현 방안 1>',
                                        '<SECTION3-A 3-1>': '<구현 방안 2>',
                                        '<SECTION3-A 6-1>': '<주요 요소 1>',
                                        '<SECTION3-A 7-1>': '<주요 요소 2>',
                                        '<SECTION3-A 10-1>': '<기대 효과 1>',
                                        '<SECTION3-A 11-1>': '<기대 효과 2>',
                                        '<SECTION3-A 13-1>': '<추가 효과 1>',
                                        '<SECTION3-A 14-1>': '<추가 효과 2>',
                                        '<SECTION5-A 2-1>': '<향후 계획 1>',
                                        '<SECTION5-A 3-1>': '<향후 계획 2>'
                                    }
                                    mapped_text = text_mapping.get(tag_value, "<제목을 입력해주세요.>")
                                    run = paragraph.add_run(mapped_text)
                                    
                                font = run.font
                                font.size = Pt(10)  # 글씨 크기
                                font.name = '맑은 고딕'  # 글씨체 설정
                                font.color.rgb = RGBColor(255, 255, 255)  # 흰색 텍스트
                                run.bold = True
                                    
            # 내용이 없는 표 제거 시도
            # print('here 삭제 시도')
            # self.clean_empty_rows(document, start_page = 12)

            document.save(output_path)
            return output_path, final_texts
        except Exception as e:
            print(f"DOCX 파일 생성 오류: {e}")
            return None

    def clean_empty_rows(self, document, start_page=None):
        """
        특정 페이지 이후의 표에서 모든 셀이 빈칸인 행을 삭제하는 함수
        
        Parameters:
        document (Document): python-docx Document 객체
        start_page (int): 검사를 시작할 페이지 번호 (None이면 모든 페이지 검사)
        """
        current_page = 1
        
        for table in document.tables:
            # 현재 테이블이 있는 페이지 확인
            if start_page is not None:
                # table._element의 부모 요소들을 검사하여 페이지 번호 찾기
                parent = table._element.getparent()
                while parent is not None:
                    if parent.tag.endswith('sectPr'):
                        current_page += 1
                        break
                    parent = parent.getparent()
                    
                if current_page < start_page:
                    continue
            
            # 삭제할 행 인덱스를 저장할 리스트
            rows_to_delete = []
            
            # 모든 행 검사
            for i, row in enumerate(table.rows):
                # 행의 모든 셀이 빈칸인지 확인
                if all(cell.text.strip() == '' for cell in row.cells):
                    rows_to_delete.append(i)
            
            # 뒤에서부터 행 삭제 (인덱스 변화 방지)
            for row_idx in reversed(rows_to_delete):
                tr = table._element.tr_lst[row_idx]
                tr.getparent().remove(tr)
                
    async def convert_docx_to_pdf(self, docx_file_path):
        pdf_filename = os.path.splitext(os.path.basename(docx_file_path))[0] + '.pdf'
        pdf_file_path = os.path.join(self.UPLOAD_FOLDER, pdf_filename)
        try:
            pythoncom.CoInitialize()
            await asyncio.to_thread(convert, docx_file_path, pdf_file_path)
            pythoncom.CoUninitialize()
            return pdf_file_path
        except Exception as e:
            print(f"DOCX->PDF 변환 오류: {e}")
            return None

    async def process_element(self, element, parent, style_mapping):
        element_type = element['type']
        try:
            if element_type == 'text':
                text_content = element.get('raw', element.get('text', ''))
                if isinstance(parent, Paragraph):
                    run = parent.add_run(text_content)
                elif isinstance(parent, Run):
                    parent.text += text_content
                else:
                    # 다른 유형의 부모가 필요한 경우 추가 처리
                    pass
            elif element_type == 'paragraph':
                if hasattr(parent, 'add_paragraph'):
                    p = parent.add_paragraph()
                    p.style = parent.style
                else:
                    p = self.insert_paragraph_before(parent)
                    
                run = p.add_run("- ")
                font = run.font
                font.name = '맑은 고딕'
                font.size = Pt(11)
                
                for child in element.get('children', []):
                    await self.process_element(child, p, style_mapping)
            
            elif element_type == 'heading':
                level = element.get('attrs', {}).get('level', 1)

                p = self.insert_paragraph_before(parent)
                
                run = p.add_run("○ ")
                
                # 제목 스타일과 내용을 설정
                # run = p.add_run()
                for child in element.get('children', []):
                    if child['type'] == 'text':
                        text_content = child.get('raw', child.get('text', ''))
                        run.add_text(text_content)
                    else:
                        await self.process_element(child, run, style_mapping)
                
                # 스타일을 설정하여 제목처럼 보이도록 합니다.
                font = run.font
                font.name = '맑은 고딕'
                font.bold = True
                font.size = Pt(14)

            # 여기 수정
            
            elif element_type in ['strong', 'underline', 'emphasis', 'strikethrough']:
                if isinstance(parent, Paragraph):
                    run = parent.add_run()
                elif isinstance(parent, Run):
                    # 동일한 단락에 새로운 Run 추가하여 스타일 적용
                    run = parent._parent.add_run()
                else:
                    # 다른 유형의 부모가 필요한 경우 추가 처리
                    pass

                # 적절한 스타일 적용
                if element_type == 'strong':
                    run.bold = True
                elif element_type == 'underline':
                    run.underline = True
                elif element_type == 'emphasis':
                    run.italic = True
                elif element_type == 'strikethrough':
                    run.font.strike = True

                for child in element.get('children', []):
                    await self.process_element(child, run, style_mapping)
            
            elif element_type == 'list':
                for item in element.get('children', []):
                    await self.process_element(item, parent, style_mapping)
                    
            elif element_type == 'list_item':
                bullet_symbol = '◦'  # 원하는 기호로 변경 가능
                if hasattr(parent, 'add_paragraph'):
                    p = parent.add_paragraph()
                    p.style = parent.style
                else:
                    p = self.insert_paragraph_before(parent)
                run = p.add_run(f"{bullet_symbol} ")
                for child in element.get('children', []):
                    await self.process_element(child, p, style_mapping)
            
            elif element_type == 'linebreak':
                parent.add_run().add_break()

            elif element_type == 'blank_line':
                parent.add_paragraph()
                
            elif element_type == 'thematic_break':
                pass
            elif element_type == 'block_text':
                for child in element.get('children', []):
                    await self.process_element(child, parent, style_mapping)
            else:
                print(f"처리되지 않은 요소 유형: {element_type}")
        except Exception as e:
            print(f"예상치 못한 오류 발생: {e}")
            if isinstance(parent, (Paragraph, Run)):
                parent.add_run(str(element))
            else:
                print(f"이 유형의 부모에 콘텐츠를 추가할 수 없습니다: {type(parent)}")

    def insert_paragraph_before(self, paragraph):
        new_p = OxmlElement('w:p')
        paragraph._element.addprevious(new_p)
        return Paragraph(new_p, paragraph._parent)

    def generate_dynamic_tags(self, keys):
        sub_tags = {}
        for key in keys:
            modified_key = re.sub(r'(SECTION\d+)-([A-Z])(\d+)', r'\1-\2 \3', key) + '-1'
            value = f"<{modified_key}>"
            sub_tags[key] = value
        return sub_tags
    
    async def parse_markdown_and_insert(self, markdown_text, parent, style_mapping):
        max_retries = 3
        attempt = 0

        while attempt < max_retries:
            attempt += 1

            response = await self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            f"사용자 제공 입력값 : {markdown_text} \n\n"
                            "다음은 사용자가 제공한 입력값에 대해 JSON 데이터를 생성하는 작업입니다. \n\n"
                            "사용자의 입력값을 분석하여, 아래 조건에 맞는 JSON 데이터를 반환하세요:\n\n"
                            "1. JSON은 {'document': ...} 형식으로 시작해야 합니다.\n"
                            "2. 입력 텍스트를 기반으로 적절한 'type'과 'children' 값을 채우세요.\n"
                            "3. 결과값은 반드시 유효한 JSON 형식이어야 합니다.\n"
                            "4. '○','-' 같은 문자가 존재한다면 그대로 포함해서 출력해야 합니다.\n"
                            "5. '○'는 제목,'-' 는 하위 내용과 관련된 기호이니 고려해야 합니다.\n"
                            "6. 사용자 제공 입력값에서 내용을 추가하거나 바꾸지 않고 그대로 사용하세요.\n"
                            "7. 입력값에 제목(○)이나 본문(-) 중 하나만 존재하는 경우, 입력값에 포함된 항목만 출력하세요.\n"
                            "   - 제목(○)만 있는 경우, 제목만 JSON 구조에 포함하세요.\n"
                            "   - 본문(-)만 있는 경우, 본문만 JSON 구조에 포함하세요.\n"
                            "   - 입력값에 없는 항목(제목 또는 본문)을 추가하거나 임의로 생성하지 마세요.\n"
                            "8. 응답은 한글로 작성하세요."
                        ),
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "import json\n\nmessages = [\n    {\n        \"role\": \"user\",\n        \"content\": json.dumps({\n            \"document\": [\n                {\n                    \"type\": \"heading\",\n                    \"attrs\": {\n                        \"level\": 1\n                    },\n                    \"children\": [\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{text1}\"\n                        }\n                    ]\n                },\n                {\n                    \"type\": \"paragraph\",\n                    \"children\": [\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{text2} \"\n                        },\n                        {\n                            \"type\": \"strong\",\n                            \"children\": [\n                                {\n                                    \"type\": \"text\",\n                                    \"raw\": \"{text3}\"\n                                }\n                            ]\n                        },\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{text4} \"\n                        },\n                        {\n                            \"type\": \"underline\",\n                            \"children\": [\n                                {\n                                    \"type\": \"text\",\n                                    \"raw\": \"{text5}\"\n                                }\n                            ]\n                        },\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{{text6} \"\n                        },\n                        {\n                            \"type\": \"underline\",\n                            \"children\": [\n                                {\n                                    \"type\": \"text\",\n                                    \"raw\": \"{text7}\"\n                                }\n                            ]\n                        },\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{text8}\"\n                        }\n                    ]\n                },\n                {\n                    \"type\": \"paragraph\",\n                    \"children\": [\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{text9} \"\n                        },\n                        {\n                            \"type\": \"strong\",\n                            \"children\": [\n                                {\n                                    \"type\": \"text\",\n                                    \"raw\": \"{text10}\"\n                                }\n                            ]\n                        },\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{text11}\"\n                        }\n                    ]\n                },\n                {\n                    \"type\": \"list\",\n                    \"children\": [\n                        {\n                            \"type\": \"list_item\",\n                            \"attrs\": {\n                                \"ordered\": false\n                            },\n                            \"children\": [\n                                {\n                                    \"type\": \"text\",\n                                    \"raw\": \"{text12}\"\n                                }\n                            ]\n                        },\n                        {\n                            \"type\": \"list_item\",\n                            \"attrs\": {\n                                \"ordered\": false\n                            },\n                            \"children\": [\n                                {\n                                    \"type\": \"text\",\n                                    \"raw\": \"{text13}\"\n                                }\n                            ]\n                        },\n                        {\n                            \"type\": \"list_item\",\n                            \"attrs\": {\n                                \"ordered\": false\n                            },\n                            \"children\": [\n                                {\n                                    \"type\": \"text\",\n                                    \"raw\": \"{text14}\"\n                                }\n                            ]\n                        }\n                    ]\n                },\n                {\n                    \"type\": \"paragraph\",\n                    \"children\": [\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{text15}\"\n                        }\n                    ]\n                },\n                {\n                    \"type\": \"thematic_break\"\n                },\n                {\n                    \"type\": \"paragraph\",\n                    \"children\": [\n                        {\n                            \"type\": \"text\",\n                            \"raw\": \"{text16}\"\n                        }\n                    ]\n                }\n            ]\n        })\n    }\n]\n\nresponse_format = \"json_object\"\n"
                            }
                        ]
                    }
                ],
                temperature=1,
                max_tokens=4096,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0,
                response_format={
                    "type": "json_object"
                }
            )

            # JSON 응답을 파싱하여 딕셔너리 형태로 `ast`에 저장
            if response:
                ast_str = response.choices[0].message.content.strip()

            # JSON 데이터를 딕셔너리로 변환
            try:
                ast = json.loads(ast_str)
            except json.JSONDecodeError as e:
                print(f'JSON 디코딩 오류: {e}')
                continue  # 재시도

            # 'document' 키가 존재하는지 확인
            if 'document' not in ast:
                print("'document' 키가 응답에 없습니다. 재시도합니다.")
                continue  # 재시도

            # false를 False로 변환하는 함수
            def convert_false_to_False(obj):
                if isinstance(obj, list):
                    return [convert_false_to_False(item) for item in obj]
                elif isinstance(obj, dict):
                    return {k: convert_false_to_False(v) if v != "false" else False for k, v in obj.items()}
                else:
                    return obj

            # JSON 객체에서 "ordered": false를 "ordered": False로 변환
            ast = convert_false_to_False(ast)

            # AST 출력 (디버깅 용도)

            try:
                # AST를 순회하며 내용 삽입
                for element in ast['document']:
                    await self.process_element(element, parent, style_mapping)
            except Exception as e:
                print(f'for문 오류 : {e}')
                # 필요한 경우 여기서 추가 예외 처리를 할 수 있습니다.

            # 성공적으로 처리되었으므로 함수 종료
            return

        # 최대 재시도 횟수를 초과한 경우 오류 처리
        print(f"최대 재시도 횟수({max_retries})를 초과했습니다.")
        # 필요한 경우 예외를 발생시키거나 다른 처리를 할 수 있습니다.
